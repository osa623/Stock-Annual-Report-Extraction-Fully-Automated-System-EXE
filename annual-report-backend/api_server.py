"""
Flask API Server for Financial Data Extractor
PDF upload → extraction with progress streaming → preview → export
"""

from flask import Flask, jsonify, request, send_file, Response
from flask_cors import CORS
import os
import sys
import json
import uuid
import logging
import tempfile
import threading
import queue
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add parent directory to path
sys.path.append(str(Path(__file__).parent))

from src.services.gemini_service import GeminiFinancialExtractor, ALL_STATEMENT_CONFIGS

# ── Flask Setup ───────────────────────────────────────────────────────
app = Flask(__name__)
CORS(app)

# ── Storage ───────────────────────────────────────────────────────────
UPLOAD_DIR = Path(__file__).parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# In-memory store
extraction_results: dict = {}
pdf_registry: dict = {}          # pdf_id -> {path, filename, ...}
progress_queues: dict = {}       # pdf_id -> Queue for SSE progress

# Initialize Gemini extractor
try:
    gemini_extractor = GeminiFinancialExtractor()
    logger.info("Extraction engine initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize extraction engine: {e}")
    gemini_extractor = None


# ══════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "ok",
        "service": "FD Extractor API",
        "engine_ready": gemini_extractor is not None,
    }), 200


@app.route('/api/upload-pdf', methods=['POST'])
def upload_pdf():
    """Accept a PDF file upload. Returns a pdf_id."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided. Use 'file' field in multipart form."}), 400

    file = request.files['file']

    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    if not file.filename.lower().endswith('.pdf'):
        return jsonify({"error": "Only PDF files are accepted"}), 400

    pdf_id = str(uuid.uuid4())[:12]
    safe_name = f"{pdf_id}_{file.filename}"
    save_path = UPLOAD_DIR / safe_name
    file.save(str(save_path))

    file_size_mb = save_path.stat().st_size / (1024 * 1024)

    pdf_registry[pdf_id] = {
        "path": str(save_path),
        "filename": file.filename,
        "uploaded_at": datetime.now().isoformat(),
        "size_mb": round(file_size_mb, 2),
    }

    logger.info(f"PDF uploaded: {file.filename} → {pdf_id} ({file_size_mb:.1f} MB)")

    return jsonify({
        "success": True,
        "pdf_id": pdf_id,
        "filename": file.filename,
        "size_mb": round(file_size_mb, 2),
    }), 200


@app.route('/api/extract/<pdf_id>', methods=['POST'])
def extract_pdf(pdf_id):
    """
    Start extraction on the uploaded PDF.
    Returns the structured data once complete.
    Also pushes progress via the SSE /api/extract/<pdf_id>/progress endpoint.
    """
    if gemini_extractor is None:
        return jsonify({"error": "Extraction engine is not available. Check server configuration."}), 503

    pdf_info = pdf_registry.get(pdf_id)
    if not pdf_info:
        return jsonify({"error": f"PDF '{pdf_id}' not found. Upload a PDF first."}), 404

    pdf_path = pdf_info["path"]
    if not Path(pdf_path).exists():
        return jsonify({"error": "PDF file no longer exists on disk"}), 404

    # Reuse queue if SSE endpoint already created one; otherwise create new
    if pdf_id not in progress_queues:
        progress_queues[pdf_id] = queue.Queue()
    q = progress_queues[pdf_id]

    def progress_cb(step, total, message):
        q.put({"step": step, "total": total, "message": message})

    try:
        logger.info(f"Starting extraction for {pdf_id}: {pdf_info['filename']}")
        result = gemini_extractor.extract_from_pdf(pdf_path, progress_callback=progress_cb)

        # Signal completion via queue
        q.put({"step": -1, "total": -1, "message": "done"})

        extraction_results[pdf_id] = {
            "data": result,
            "filename": pdf_info["filename"],
            "extracted_at": datetime.now().isoformat(),
        }

        return jsonify({
            "success": True,
            "pdf_id": pdf_id,
            "filename": pdf_info["filename"],
            "data": result,
        }), 200

    except FileNotFoundError as e:
        q.put({"step": -1, "total": -1, "message": f"error:{e}"})
        return jsonify({"error": str(e)}), 404
    except ValueError as e:
        q.put({"step": -1, "total": -1, "message": f"error:{e}"})
        return jsonify({"error": str(e)}), 400
    except RuntimeError as e:
        q.put({"step": -1, "total": -1, "message": f"error:{e}"})
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        q.put({"step": -1, "total": -1, "message": f"error:{e}"})
        logger.error(f"Extraction error: {e}", exc_info=True)
        return jsonify({"error": f"Extraction failed: {str(e)}"}), 500
    finally:
        progress_queues.pop(pdf_id, None)


@app.route('/api/extract/<pdf_id>/<statement_key>', methods=['POST'])
def extract_single_statement(pdf_id, statement_key):
    """
    Extract a SINGLE financial statement from the uploaded PDF.
    Returns just that statement's data. Results accumulate in extraction_results.
    """
    if gemini_extractor is None:
        return jsonify({"error": "Extraction engine is not available. Check server configuration."}), 503

    if statement_key not in ALL_STATEMENT_CONFIGS:
        return jsonify({
            "error": f"Unknown statement type: '{statement_key}'",
            "valid_keys": list(ALL_STATEMENT_CONFIGS.keys()),
        }), 400

    pdf_info = pdf_registry.get(pdf_id)
    if not pdf_info:
        return jsonify({"error": f"PDF '{pdf_id}' not found. Upload a PDF first."}), 404

    pdf_path = pdf_info["path"]
    if not Path(pdf_path).exists():
        return jsonify({"error": "PDF file no longer exists on disk"}), 404

    display_name = ALL_STATEMENT_CONFIGS[statement_key][1]

    try:
        logger.info(f"Single extraction [{statement_key}] for {pdf_id}: {pdf_info['filename']}")
        section_data = gemini_extractor.extract_single(pdf_path, statement_key)

        # Accumulate results for this PDF
        if pdf_id not in extraction_results:
            extraction_results[pdf_id] = {
                "data": {},
                "filename": pdf_info["filename"],
                "extracted_at": datetime.now().isoformat(),
            }
        extraction_results[pdf_id]["data"][statement_key] = section_data
        extraction_results[pdf_id]["extracted_at"] = datetime.now().isoformat()

        return jsonify({
            "success": True,
            "pdf_id": pdf_id,
            "statement_key": statement_key,
            "display_name": display_name,
            "data": section_data,
        }), 200

    except FileNotFoundError as e:
        return jsonify({"error": str(e)}), 404
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        logger.error(f"Single extraction error [{statement_key}]: {e}", exc_info=True)
        return jsonify({"error": f"Extraction failed: {str(e)}"}), 500


@app.route('/api/statements', methods=['GET'])
def list_statement_types():
    """Return all available statement types for extraction."""
    types = []
    for key, (prompt, display_name) in ALL_STATEMENT_CONFIGS.items():
        types.append({"key": key, "display_name": display_name})
    return jsonify({"statement_types": types}), 200


@app.route('/api/extract/<pdf_id>/progress', methods=['GET'])
def extraction_progress(pdf_id):
    """
    Server-Sent Events endpoint for real-time extraction progress.
    Frontend connects to this BEFORE calling POST /api/extract/<pdf_id>.
    """
    # Create the queue if it doesn't exist yet (frontend connects first)
    if pdf_id not in progress_queues:
        progress_queues[pdf_id] = queue.Queue()

    q = progress_queues[pdf_id]

    def generate():
        while True:
            try:
                msg = q.get(timeout=120)  # 2 min timeout
                yield f"data: {json.dumps(msg)}\n\n"
                if msg.get("step") == -1:
                    break
            except queue.Empty:
                # Send keepalive
                yield f"data: {json.dumps({'step': 0, 'total': 0, 'message': 'waiting...'})}\n\n"

    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive',
        }
    )


@app.route('/api/export/<pdf_id>', methods=['POST'])
def export_data(pdf_id):
    """Export extracted data in the requested format."""
    export_format = request.args.get('format', 'json').lower()
    valid_formats = {'json', 'xlsx', 'csv', 'pdf', 'docx'}

    if export_format not in valid_formats:
        return jsonify({"error": f"Invalid format '{export_format}'. Use: {', '.join(valid_formats)}"}), 400

    stored = extraction_results.get(pdf_id)
    if not stored:
        return jsonify({"error": f"No extraction results for '{pdf_id}'. Run extraction first."}), 404

    data = stored["data"]
    base_name = Path(stored["filename"]).stem

    try:
        if export_format == 'json':
            return _export_json(data, base_name)
        elif export_format == 'xlsx':
            return _export_xlsx(data, base_name)
        elif export_format == 'csv':
            return _export_csv(data, base_name)
        elif export_format == 'pdf':
            return _export_pdf(data, base_name)
        elif export_format == 'docx':
            return _export_docx(data, base_name)
    except Exception as e:
        logger.error(f"Export error ({export_format}): {e}", exc_info=True)
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


# ══════════════════════════════════════════════════════════════════════
# EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════

def _build_section_list(data):
    """Build a unified section list from extraction data (supports old and new keys)."""
    sections = [
        ("Income Statement", data.get("income_statement")),
        ("Statement of Financial Position", data.get("balance_sheet")),
        ("Cash Flow Statement", data.get("cash_flow")),
        ("Statement of Comprehensive Income", data.get("comprehensive_income")),
        ("Statement of Changes in Equity", data.get("changes_in_equity")),
        ("Independent Auditor's Report", data.get("auditors_report")),
    ]
    # Also include legacy additional_sections
    for i, sec in enumerate(data.get("additional_sections") or []):
        title = sec.get("title", f"Additional {i+1}")
        sections.append((title, sec))
    return sections


def _section_to_rows(section):
    if not section or not isinstance(section, dict):
        return [], []
    headers = section.get("headers", [])
    rows = []
    for row in section.get("rows", []):
        row_data = [row.get("item", "")]
        values = row.get("values", [])
        row_data.extend(values)
        rows.append(row_data)
    return headers, rows


def _export_json(data, base_name):
    tmp = tempfile.NamedTemporaryFile(
        suffix=".json", prefix=f"{base_name}_", delete=False, mode='w', encoding='utf-8'
    )
    json.dump(data, tmp, indent=2, ensure_ascii=False)
    tmp.close()
    return send_file(tmp.name, as_attachment=True, download_name=f"{base_name}_extracted.json", mimetype="application/json")


def _export_xlsx(data, base_name):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    sections = _build_section_list(data)

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    for sheet_name, section in sections:
        if not section:
            continue
        headers, rows = _section_to_rows(section)
        if not rows:
            continue

        ws = wb.create_sheet(title=sheet_name[:31])
        title_text = section.get("title", sheet_name)
        ws.append([title_text])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        ws.cell(row=1, column=1).font = Font(bold=True, size=14)
        ws.append([])

        notes = section.get("notes", "")
        if notes:
            ws.append([f"Note: {notes}"])
            ws.append([])

        header_row_idx = ws.max_row + 1
        ws.append(headers)
        for col_idx, _ in enumerate(headers, 1):
            cell = ws.cell(row=header_row_idx, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        for row_data in rows:
            ws.append(row_data)
            for col_idx in range(1, len(row_data) + 1):
                cell = ws.cell(row=ws.max_row, column=col_idx)
                cell.border = thin_border
                if col_idx > 1:
                    cell.alignment = Alignment(horizontal="right")
                    if cell.value is not None and isinstance(cell.value, (int, float)):
                        cell.number_format = '#,##0'

        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

    if not wb.sheetnames:
        ws = wb.create_sheet("No Data")
        ws.append(["No financial statements were extracted."])

    tmp_path = Path(tempfile.gettempdir()) / f"{base_name}_extracted.xlsx"
    wb.save(str(tmp_path))
    return send_file(str(tmp_path), as_attachment=True, download_name=f"{base_name}_extracted.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def _export_csv(data, base_name):
    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output)

    sections = _build_section_list(data)

    for section_name, section in sections:
        if not section:
            continue
        headers, rows = _section_to_rows(section)
        if not rows:
            continue
        writer.writerow([])
        writer.writerow([f"=== {section_name} ==="])
        writer.writerow(headers)
        for row_data in rows:
            writer.writerow(row_data)

    tmp = tempfile.NamedTemporaryFile(
        suffix=".csv", prefix=f"{base_name}_", delete=False, mode='w', encoding='utf-8-sig', newline=''
    )
    tmp.write(output.getvalue())
    tmp.close()
    return send_file(tmp.name, as_attachment=True, download_name=f"{base_name}_extracted.csv", mimetype="text/csv")


def _export_pdf(data, base_name):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    tmp_path = Path(tempfile.gettempdir()) / f"{base_name}_extracted.pdf"
    doc = SimpleDocTemplate(str(tmp_path), pagesize=landscape(A4), topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
    section_style = ParagraphStyle('SectionTitle', parent=styles['Heading2'], fontSize=14,
        textColor=colors.HexColor("#2F5496"), spaceAfter=8, spaceBefore=16)

    elements.append(Paragraph(f"Financial Data Report — {base_name}", title_style))
    elements.append(Paragraph(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))

    sections = _build_section_list(data)

    for section_name, section in sections:
        if not section:
            continue
        headers, rows = _section_to_rows(section)
        if not rows:
            continue

        elements.append(Paragraph(section_name, section_style))
        table_data = [headers] + rows
        col_count = len(headers) if headers else (len(rows[0]) if rows else 1)
        available_width = landscape(A4)[0] - 60
        first_col = available_width * 0.4
        other_col = (available_width * 0.6) / max(col_count - 1, 1) if col_count > 1 else available_width
        col_widths = [first_col] + [other_col] * (col_count - 1)

        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2F5496")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

    doc.build(elements)
    return send_file(str(tmp_path), as_attachment=True, download_name=f"{base_name}_extracted.pdf", mimetype="application/pdf")


def _export_docx(data, base_name):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    doc.add_heading(f"Financial Data Report — {base_name}", level=0)
    doc.add_paragraph(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    sections = _build_section_list(data)

    for section_name, section in sections:
        if not section:
            continue
        headers, rows = _section_to_rows(section)
        if not rows:
            continue

        doc.add_heading(section_name, level=2)
        notes = section.get("notes", "")
        if notes:
            p = doc.add_paragraph(f"Note: {notes}")
            p.runs[0].italic = True

        col_count = len(headers) if headers else (len(rows[0]) if rows else 1)
        table = doc.add_table(rows=1 + len(rows), cols=col_count, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(h) if h else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, val in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(val) if val is not None else ""
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
        doc.add_paragraph("")

    tmp_path = Path(tempfile.gettempdir()) / f"{base_name}_extracted.docx"
    doc.save(str(tmp_path))
    return send_file(str(tmp_path), as_attachment=True, download_name=f"{base_name}_extracted.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    logger.info("Starting FD Extractor API Server...")
    logger.info(f"Upload directory: {UPLOAD_DIR}")

    app.run(
        host='0.0.0.0',
        port=5000,
        debug=True,
        threaded=True,
    )
